from __future__ import annotations

import base64
import html
import json
import mimetypes
import re
import uuid
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any, Callable
from urllib.parse import parse_qs, urlparse


ROOT = Path(__file__).resolve().parent.parent


TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".log", ".yaml", ".yml"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg", ".avif", ".tif", ".tiff"}
SPREADSHEET_EXTENSIONS = {".xlsx", ".xlsm", ".xltx", ".xltm", ".xls"}


def json_response(handler: BaseHTTPRequestHandler, payload: Any, status: int = 200):
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Content-Length", str(len(body)))
    handler.send_header("Access-Control-Allow-Origin", "*")
    handler.send_header("Access-Control-Allow-Methods", "GET,POST,PATCH,OPTIONS")
    handler.send_header("Access-Control-Allow-Headers", "Content-Type")
    handler.end_headers()
    handler.wfile.write(body)


def file_response(
    handler: BaseHTTPRequestHandler,
    path: Path,
    download_name: str,
    content_type: str | None = None,
    inline: bool = False,
):
    if not path.exists():
        json_response(handler, {"error": "File not found"}, 404)
        return
    data = path.read_bytes()
    handler.send_response(200)
    handler.send_header("Content-Type", content_type or mimetypes.guess_type(download_name)[0] or "application/octet-stream")
    disposition = "inline" if inline else "attachment"
    handler.send_header("Content-Disposition", f'{disposition}; filename="{download_name}"')
    handler.send_header("Content-Length", str(len(data)))
    handler.send_header("Access-Control-Allow-Origin", "*")
    handler.end_headers()
    handler.wfile.write(data)


def _read_pdf_text(file_path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        return f"PDF text preview requires pypdf. {exc}"

    reader = PdfReader(str(file_path))
    pages = []
    for index, page in enumerate(reader.pages, 1):
        pages.append(f"--- Page {index} ---\n{page.extract_text() or ''}".strip())
    return "\n\n".join(pages)


def _read_docx_html(file_path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:
        return f"<p>DOCX preview requires python-docx. {html.escape(str(exc))}</p>"

    document = Document(str(file_path))
    parts = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(f"<p>{html.escape(text)}</p>")
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = "".join(f"<td>{html.escape(cell.text.strip())}</td>" for cell in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        parts.append(f"<table>{''.join(rows)}</table>")
    return "\n".join(parts) or "<p>No readable text found in this document.</p>"


def _read_workbook_html(file_path: Path) -> str:
    if file_path.suffix.lower() == ".xls":
        try:
            import xlrd
        except Exception as exc:
            return f"<p>Legacy XLS preview requires xlrd. {html.escape(str(exc))}</p>"

        workbook = xlrd.open_workbook(str(file_path))
        sheets = []
        for sheet in workbook.sheets():
            rows = []
            for row_index in range(sheet.nrows):
                cells = "".join(
                    f"<td>{html.escape(str(sheet.cell_value(row_index, column_index)))}</td>"
                    for column_index in range(sheet.ncols)
                )
                rows.append(f"<tr>{cells}</tr>")
            sheets.append(
                f"<section><h2>{html.escape(sheet.name)}</h2><table>{''.join(rows)}</table></section>"
            )
        return "\n".join(sheets) or "<p>No readable sheets found in this workbook.</p>"

    try:
        from openpyxl import load_workbook
    except Exception as exc:
        return f"<p>Workbook preview requires openpyxl. {html.escape(str(exc))}</p>"

    workbook = load_workbook(str(file_path), data_only=True, read_only=True)
    sheets = []
    for sheet in workbook.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            cells = "".join(
                f"<td>{html.escape('' if value is None else str(value))}</td>"
                for value in row
            )
            rows.append(f"<tr>{cells}</tr>")
        sheets.append(
            f"<section><h2>{html.escape(sheet.title)}</h2><table>{''.join(rows)}</table></section>"
        )
    workbook.close()
    return "\n".join(sheets) or "<p>No readable sheets found in this workbook.</p>"


def read_json_body(handler: BaseHTTPRequestHandler) -> dict:
    length = int(handler.headers.get("Content-Length", "0") or "0")
    if length <= 0:
        return {}
    return json.loads(handler.rfile.read(length).decode("utf-8-sig"))


def safe_file_name(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", value).strip()
    return cleaned or "uploaded_file"


def parse_multipart(handler: BaseHTTPRequestHandler) -> tuple[dict[str, str], list[dict[str, Any]]]:
    content_type = handler.headers.get("Content-Type", "")
    match = re.search(r"boundary=(.+)", content_type)
    if not match:
        return {}, []
    boundary = match.group(1).strip('"').encode()
    length = int(handler.headers.get("Content-Length", "0") or "0")
    body = handler.rfile.read(length)
    fields: dict[str, str] = {}
    files: list[dict[str, Any]] = []
    for raw_part in body.split(b"--" + boundary):
        part = raw_part.strip(b"\r\n")
        if not part or part == b"--" or b"\r\n\r\n" not in part:
            continue
        raw_headers, content = part.split(b"\r\n\r\n", 1)
        header_text = raw_headers.decode("utf-8", errors="replace")
        name_match = re.search(r'name="([^"]+)"', header_text)
        if not name_match:
            continue
        name = name_match.group(1)
        filename_match = re.search(r'filename="([^"]*)"', header_text)
        if filename_match:
            filename = safe_file_name(filename_match.group(1))
            content_type_match = re.search(r"Content-Type:\s*([^\r\n]+)", header_text, re.IGNORECASE)
            files.append(
                {
                    "field": name,
                    "filename": filename,
                    "content_type": content_type_match.group(1).strip() if content_type_match else "application/octet-stream",
                    "content": content.rstrip(b"\r\n"),
                }
            )
        else:
            fields[name] = content.decode("utf-8", errors="replace").rstrip("\r\n")
    return fields, files


def read_display_content(file_path: Path, fallback_name: str = "") -> dict:
    file_name = fallback_name or file_path.name
    extension = file_path.suffix.lower() or Path(file_name).suffix.lower()
    if not file_path.exists():
        return {
            "fileName": file_name,
            "extension": extension,
            "sizeBytes": 0,
            "mode": "missing",
            "message": "The file path recorded in metadata no longer exists.",
        }
    size = file_path.stat().st_size
    if extension in IMAGE_EXTENSIONS:
        mime = "image/svg+xml" if extension == ".svg" else f"image/{extension.replace('.', '')}"
        return {
            "fileName": file_name,
            "extension": extension,
            "sizeBytes": size,
            "mode": "image",
            "imageDataUrl": f"data:{mime};base64,{base64.b64encode(file_path.read_bytes()).decode('ascii')}",
        }
    if extension in TEXT_EXTENSIONS:
        return {
            "fileName": file_name,
            "extension": extension,
            "sizeBytes": size,
            "mode": "text",
            "text": file_path.read_text(encoding="utf-8-sig", errors="replace"),
        }
    if extension == ".pdf":
        return {
            "fileName": file_name,
            "extension": extension,
            "sizeBytes": size,
            "mode": "pdf",
            "text": _read_pdf_text(file_path),
        }
    if extension == ".docx":
        return {
            "fileName": file_name,
            "extension": extension,
            "sizeBytes": size,
            "mode": "office",
            "html": _read_docx_html(file_path),
        }
    if extension in SPREADSHEET_EXTENSIONS:
        return {
            "fileName": file_name,
            "extension": extension,
            "sizeBytes": size,
            "mode": "office",
            "html": _read_workbook_html(file_path),
        }
    return {
        "fileName": file_name,
        "extension": extension,
        "sizeBytes": size,
        "mode": "binary",
        "message": "Preview is not available for this binary file type. Use download to inspect the original.",
    }


class JsonApiHandler(BaseHTTPRequestHandler):
    routes: dict[str, Callable[[BaseHTTPRequestHandler], None]] = {}

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET,POST,PATCH,OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_GET(self):
        self._dispatch("GET")

    def do_POST(self):
        self._dispatch("POST")

    def do_PATCH(self):
        self._dispatch("PATCH")

    def _dispatch(self, method: str):
        parsed = urlparse(self.path)
        key = f"{method} {parsed.path}"
        handler = self.routes.get(key)
        if handler:
            try:
                handler(self)
            except Exception as exc:
                json_response(self, {"error": str(exc)}, 500)
            return
        for route_key, route_handler in self.routes.items():
            route_method, route_pattern = route_key.split(" ", 1)
            if route_method != method:
                continue
            pattern = "^" + re.sub(r"<([^/]+)>", r"(?P<\1>[^/]+)", route_pattern) + "$"
            match = re.match(pattern, parsed.path)
            if match:
                self.path_params = match.groupdict()  # type: ignore[attr-defined]
                try:
                    route_handler(self)
                except Exception as exc:
                    json_response(self, {"error": str(exc)}, 500)
                return
        json_response(self, {"error": "Not found"}, 404)

    def log_message(self, format: str, *args):
        return


def query_params(handler: BaseHTTPRequestHandler) -> dict[str, str]:
    parsed = urlparse(handler.path)
    values = parse_qs(parsed.query)
    return {key: item[0] for key, item in values.items() if item}


def run_server(handler_class: type[JsonApiHandler], port: int, name: str):
    server = ThreadingHTTPServer(("0.0.0.0", port), handler_class)
    print(f"{name} API running on http://localhost:{port}", flush=True)
    server.serve_forever()
