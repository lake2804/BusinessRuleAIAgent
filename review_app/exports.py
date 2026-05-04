"""Build downloadable corrected-file artifacts from structured review results."""
from __future__ import annotations

import csv
import base64
import io
import json
from typing import Any, Dict, List, Optional

from review_app.schema import ReviewStructuredOutput


Artifact = Dict[str, str]


def _corrected_records(structured_output: ReviewStructuredOutput) -> List[Dict[str, Any]]:
    if structured_output.corrected_records:
        return structured_output.corrected_records
    return [case.corrected_record for case in structured_output.cases if case.corrected_record]


def _original_records(parsed_file: Dict[str, Any]) -> List[Dict[str, Any]]:
    records = parsed_file.get("metadata", {}).get("records") or []
    return [dict(record) for record in records if isinstance(record, dict)]


def _record_id(record: Dict[str, Any], index: int) -> str:
    return str(record.get("case_id") or record.get("id") or record.get("case") or f"row_{index + 1}")


def build_corrected_records(
    parsed_file: Optional[Dict[str, Any]],
    structured_output: Optional[ReviewStructuredOutput],
) -> List[Dict[str, Any]]:
    """Merge model corrections with original records for deterministic exports."""
    if not parsed_file or not structured_output:
        return []

    originals = _original_records(parsed_file)
    corrections = _corrected_records(structured_output)
    correction_by_id = {
        _record_id(record, index): record for index, record in enumerate(corrections)
    }
    case_by_id = {case.case_id: case for case in structured_output.cases}

    if originals:
        merged = []
        for index, original in enumerate(originals):
            case_id = _record_id(original, index)
            corrected = dict(original)
            corrected.update(correction_by_id.get(case_id, {}))
            case = case_by_id.get(case_id)
            if case:
                corrected.setdefault("business_rule_final_decision_allowed", case.final_decision_allowed)
                corrected.setdefault("business_rule_required_owner", case.required_owner_or_approver)
                corrected.setdefault("business_rule_correction_note", case.proposed_corrected_resolution)
                corrected.setdefault(
                    "business_rule_needs_confirmation",
                    case.inference_or_needs_confirmation,
                )
            merged.append(corrected)
        return merged

    return corrections


def build_corrected_file_exports(
    parsed_file: Optional[Dict[str, Any]],
    structured_output: Optional[ReviewStructuredOutput],
) -> Dict[str, Artifact]:
    """Return export artifacts keyed by format name."""
    if not parsed_file or not structured_output:
        return {}

    records = build_corrected_records(parsed_file, structured_output)
    if not records:
        return {}

    base_name = parsed_file.get("file_name", "corrected_file").rsplit(".", 1)[0]
    exports: Dict[str, Artifact] = {}

    exports["json"] = {
        "file_name": f"{base_name}_corrected.json",
        "mime": "application/json",
        "content": json.dumps(records, ensure_ascii=False, indent=2),
        "encoding": "text",
    }

    exports["txt"] = {
        "file_name": f"{base_name}_corrected.txt",
        "mime": "text/plain",
        "content": "\n\n".join(json.dumps(record, ensure_ascii=False) for record in records),
        "encoding": "text",
    }

    fieldnames = sorted({key for record in records for key in record.keys()})
    if not fieldnames:
        fieldnames = ["value"]
        records = [{"value": json.dumps(record, ensure_ascii=False)} for record in records]
    if fieldnames:
        buffer = io.StringIO()
        writer = csv.DictWriter(buffer, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(records)
        exports["csv"] = {
            "file_name": f"{base_name}_corrected.csv",
            "mime": "text/csv",
            "content": buffer.getvalue(),
            "encoding": "text",
        }
        exports["xlsx"] = _build_xlsx_export(base_name, records, fieldnames)

    lines = ["# Corrected File", ""]
    for index, record in enumerate(records, 1):
        lines.append(f"## Record {index}")
        for key in fieldnames or record.keys():
            lines.append(f"- {key}: {record.get(key, '')}")
        lines.append("")
    exports["md"] = {
        "file_name": f"{base_name}_corrected.md",
        "mime": "text/markdown",
        "content": "\n".join(lines),
        "encoding": "text",
    }
    exports["docx"] = _build_docx_export(base_name, records, fieldnames)
    exports["pdf"] = _build_pdf_export(base_name, records, fieldnames)

    return exports


def _binary_artifact(file_name: str, mime: str, data: bytes) -> Artifact:
    return {
        "file_name": file_name,
        "mime": mime,
        "content": base64.b64encode(data).decode("ascii"),
        "encoding": "base64",
    }


def _build_xlsx_export(base_name: str, records: List[Dict[str, Any]], fieldnames: List[str]) -> Artifact:
    try:
        from openpyxl import Workbook
    except ImportError as exc:
        return _missing_dependency_artifact(base_name, "xlsx", "openpyxl", exc)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Corrected"
    sheet.append(fieldnames)
    for record in records:
        sheet.append([record.get(field, "") for field in fieldnames])
    for column_cells in sheet.columns:
        max_len = max(len(str(cell.value or "")) for cell in column_cells)
        sheet.column_dimensions[column_cells[0].column_letter].width = min(max(max_len + 2, 12), 48)

    buffer = io.BytesIO()
    workbook.save(buffer)
    return _binary_artifact(
        f"{base_name}_corrected.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        buffer.getvalue(),
    )


def _build_docx_export(base_name: str, records: List[Dict[str, Any]], fieldnames: List[str]) -> Artifact:
    try:
        import docx
    except ImportError as exc:
        return _missing_dependency_artifact(base_name, "docx", "python-docx", exc)

    document = docx.Document()
    document.add_heading("Corrected File", level=1)
    table = document.add_table(rows=1, cols=max(len(fieldnames), 1))
    table.style = "Table Grid"
    for index, field in enumerate(fieldnames):
        table.rows[0].cells[index].text = field
    for record in records:
        row = table.add_row()
        for index, field in enumerate(fieldnames):
            row.cells[index].text = str(record.get(field, ""))

    buffer = io.BytesIO()
    document.save(buffer)
    return _binary_artifact(
        f"{base_name}_corrected.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buffer.getvalue(),
    )


def _build_pdf_export(base_name: str, records: List[Dict[str, Any]], fieldnames: List[str]) -> Artifact:
    try:
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
        from reportlab.lib import colors
    except ImportError as exc:
        return _missing_dependency_artifact(base_name, "pdf", "reportlab", exc)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), leftMargin=24, rightMargin=24)
    rows = [fieldnames]
    for record in records:
        rows.append([str(record.get(field, ""))[:120] for field in fieldnames])
    table = Table(rows, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8f1ff")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#9aa4b2")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    doc.build([table])
    return _binary_artifact(f"{base_name}_corrected.pdf", "application/pdf", buffer.getvalue())


def _missing_dependency_artifact(base_name: str, extension: str, package: str, exc: Exception) -> Artifact:
    return {
        "file_name": f"{base_name}_corrected_{extension}_unavailable.txt",
        "mime": "text/plain",
        "content": f"Cannot build .{extension} export because {package} is not installed: {exc}",
        "encoding": "text",
    }
